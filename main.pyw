import os
import sys
import subprocess  # 폴더 열기 위해 subprocess 사용


from docx import Document
from datetime import datetime
from tkinter import Tk, Label, Button, Frame, Text, Scrollbar, RIGHT, Y, BOTH, END
from docx.shared import Pt  # 포인트 단위 사용
from docx.oxml.ns import qn  # 글꼴 변경 시 필요
from datetime import datetime


# 전역 변수로 로그 Text 위젯을 참조할 수 있도록 설정
log_text_widget = None


def get_template_path():
    # doc_templates 폴더의 경로 반환 (패키징된 경우와 개발 중 경우 모두 지원)
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller로 패키징된 임시 폴더에서 경로를 가져옴
        return sys._MEIPASS
    return os.path.abspath(os.path.dirname(__file__))


def get_current_path():
    # 현재 실행 파일의 디렉토리를 반환 (입출력 폴더 경로 기반)
    return os.path.abspath(os.path.dirname(sys.argv[0]))

def settingUi(window):
    global log_text_widget  # 전역 변수로 설정하여 다른 함수에서 접근 가능

    window.title("법인설립서 자동작성 프로그램")  # 프로그램 타이틀

    screen_width = window.winfo_screenwidth()  # 화면 너비
    screen_height = window.winfo_screenheight()  # 화면 높이

    window.geometry("800x900+{}+{}".format(int(screen_width / 2 - 400), int(screen_height / 2 - 450)))  # 프로그램 크기 및 위치
    window.resizable(False, False)  # 프로그램 크기 조절 불가능

    # 프로그램 설명 프레임
    desc_frame = Frame(window, bg="#f0f0f0", padx=10, pady=10)
    desc_frame.pack(fill='x', pady=10)

    Label(desc_frame, text="기능 설명:", bg="#f0f0f0", anchor='w').pack(anchor='w')
    Label(desc_frame, text="- '입출력 폴더생성': 입력 및 출력 폴더를 생성합니다.", bg="#f0f0f0", anchor='w').pack(pady=2, padx=10, anchor='w')
    Label(desc_frame, text="- '법인서류 자동생성': 입력 폴더 내의 엑셀 파일을 읽고 문서를 생성합니다.", bg="#f0f0f0", anchor='w').pack(pady=2, padx=10, anchor='w')

    # 버튼 프레임
    button_frame = Frame(window)
    button_frame.pack(pady=20, fill='x')

    # 로그 프레임
    log_frame = Frame(window, padx=20, pady=15)
    log_frame.pack(pady=10, fill='both', expand=True)

    scrollbar = Scrollbar(log_frame)
    scrollbar.pack(side=RIGHT, fill=Y)

    log_text_widget = Text(log_frame, wrap='word', yscrollcommand=scrollbar.set)
    log_text_widget.pack(fill=BOTH, expand=True, padx=5, pady=5)

    scrollbar.config(command=log_text_widget.yview)

    # 버튼 추가
    Button(button_frame, text="입출력 폴더생성", width=20, height=2, command=createFolders).pack(side='left', expand=True, padx=10)
    Button(button_frame, text="법인서류 자동생성", width=20, height=2, command=checkInputExcel).pack(side='left', expand=True, padx=10)
    Button(button_frame, text="Templelet 폴더 열기", width=20, height=2, command=open_template_folder).pack(side='left', expand=True, padx=10)

def open_template_folder():
    # 템플릿 폴더 경로
    base_path = get_template_path()
    template_folder = os.path.join(base_path, 'doc_templates')

    # 템플릿 폴더가 존재하는지 확인하고 열기
    if os.path.exists(template_folder):
        if os.name == 'nt':  # Windows
            subprocess.Popen(f'explorer "{template_folder}"')
        elif os.name == 'posix':  # macOS 또는 Linux
            subprocess.Popen(['open', template_folder] if sys.platform == 'darwin' else ['xdg-open', template_folder])
    else:
        log_message("템플릿 폴더가 존재하지 않습니다", 'red')



def log_message(message, color='black'):
    if log_text_widget:
        tag_name = f"tag_{log_text_widget.index('end')}"  # 각 메시지에 고유한 태그 생성
        log_text_widget.insert(END, message + "\n", tag_name)
        log_text_widget.tag_config(tag_name, foreground=color)
        log_text_widget.see(END)

def createFolders():
    base_path = get_current_path()
    input_folder = os.path.join(base_path, '입력폴더(Excel)')
    output_folder = os.path.join(base_path, '출력폴더(Word)')


    checkFolders(input_folder, "입력")
    checkFolders(output_folder, "출력")

def checkFolders(folder_path, folder_name):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        log_message(f"{folder_name} 폴더가 생성되었습니다", 'green')
    else:
        log_message(f"{folder_name} 폴더가 이미 존재합니다", 'blue')

def checkInputExcel():
    base_path = get_current_path()
    input_folder_path = os.path.join(base_path, "입력폴더(Excel)")
    input_files = [f for f in os.listdir(input_folder_path) if f.endswith('.xlsx')]

    file_count = len(input_files)
    log_message(f"입력 폴더에 {file_count}개의 엑셀 파일이 발견되었습니다", 'blue')

    if not input_files:
        log_message("입력 폴더에 엑셀 파일이 존재하지 않습니다", 'red')
        return

    for idx, file_name in enumerate(input_files, start=1):
        log_message(f"({idx}/{file_count}) '{file_name}' 파일을 처리 중입니다...", 'blue')
        read_excel_data(os.path.join(input_folder_path, file_name))


def read_excel_data(filepath):
    from openpyxl import load_workbook
    workbook = load_workbook(filepath, data_only=True)
    corp_info, directors, shareholders, welcome_info = None, [], [], None

    try:
        corp_info = read_corpInfo(workbook)
        directors = read_directorIinfo(workbook)
        shareholders, shareholder_count = read_shareholderInfo(workbook)
        ceo_info = read_representativeInfo(workbook)
        basic_info = read_basicInfo(workbook)
        log_message(f"엑셀 파일 '{os.path.basename(filepath)}'의 정보를 불러왔습니다", 'blue')
        log_message(f"법인설립정보: {corp_info}")
        log_message(f"임원 정보: {directors}")
        log_message(f"주주 정보: {shareholders}")
        log_message(f"대표자 정보: {ceo_info}")
        log_message(f"기본 정보: {basic_info}")


        generate_word_doc(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count  ,filepath)
        generate_word_doc_1(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count  ,filepath)

    except KeyError as e:
        log_message(f"'{filepath}' 파일에서 시트를 찾을 수 없습니다: {e}", 'red')
    except Exception as e:
        log_message(f"'{filepath}' 파일 처리 중 에러 발생: {e}", 'red')
    finally:
        workbook.close()

def read_corpInfo(workbook):
    sheet = workbook['법인설립정보']
    return {
        "법인명": sheet['C3'].value,
        "영문법인명": sheet['C5'].value,
        "본점주소": sheet['C7'].value,
        "주당금액": sheet['C8'].value,
        "최대주식수": sheet['C9'].value,
        "발행주식수": sheet['C10'].value,
        "주식타입": sheet['C11'].value,
        "자본금": (sheet['C8'].value or 0) * (sheet['C10'].value or 0),
        "설립목적": sheet['C13'].value,
        "증자": sheet['C14'].value,
        "신주인수권부사채발행": sheet['C15'].value,
        "중간배당": sheet['C16'].value,
        "전환사채": sheet['C17'].value
    }

def read_directorIinfo(workbook):
    sheet = workbook['임원 주주 정보']
    directors = {}
    
    for i in range(3, 11):
        position = sheet[f'B{i}'].value
        nationality = sheet[f'C{i}'].value
        name = sheet[f'D{i}'].value
        english_name = sheet[f'E{i}'].value
        birth_date = sheet[f'F{i}'].value
        address = sheet[f'G{i}'].value

        if not position or not name:
            continue  # 직책이나 이름이 없으면 건너뜀

        # 생년월일 포맷 변경: 외국인일 경우 "Y년 M월 D일", 한국인일 경우 주민등록번호 형태로
        if birth_date:
            if isinstance(birth_date, datetime):  # birth_date가 datetime 객체인지 확인
                if any(keyword in nationality.lower() for keyword in ["한국", "korean", "korea", "한국인"]):
                    birth_date = birth_date.strftime('%Y%m%d')  # 주민등록번호 형태
                else:
                    birth_date = birth_date.strftime('%Y년 %m월 %d일')  # 외국인 포맷
            else:
                # birth_date가 문자열일 경우 그대로 사용
                birth_date = str(birth_date)

        if i == 3:
            if "한국" in nationality or "korean" in nationality.lower() or "korea" in nationality.lower() or "한국인" in nationality:
                directors['대표피선자'] = f"사내이사 {name} ({birth_date})" + '\n'
                director_info = f"{position} {name} ({birth_date}) {address}" + '\n'
            else:
                directors['대표피선자'] = f"사내이사 {name} (영문:{english_name}, 국적:{nationality}, 생년월일:{birth_date})" + '\n'
                director_info = f"{position} {nationality} {name} {english_name} ({birth_date}) {address}" + '\n'
        else:
            if "한국" in nationality or "korean" in nationality.lower() or "korea" in nationality.lower() or "한국인" in nationality:
                director_info = f"{position} {name} ({birth_date})" + '\n'
            else:
                director_info = f"{position} {nationality} {name} {english_name} ({birth_date})" + '\n'
        
        directors[f'임원{i - 2}'] = director_info + '\n'
    
    return directors


def read_shareholderInfo(workbook):
    sheet = workbook['임원 주주 정보']
    shareholders = {}
    shareholder_count = 0  # 주주 숫자 초기화

    for i in range(15, 23):
        row = [sheet[f'{chr(66 + j)}{i}'].value if sheet[f'{chr(66 + j)}{i}'].value is not None else '' for j in range(6)]
        
        # 모든 데이터가 비어 있는 경우 건너뜀
        if not any(row):
            continue
        
        shareholder_count += 1  # 주주가 발견될 때마다 증가
        shareholders[f'주주{shareholder_count}'] = row
    
    return shareholders, shareholder_count  # 주주 정보와 주주 숫자 반환



def read_representativeInfo(workbook):
    sheet = workbook['임원 주주 정보']

    return{
        "대표자직책": sheet['B3'].value,
        "대표자국적": sheet['C3'].value,
        "대표자명": sheet['D3'].value,
        "대표자영문명": sheet['E3'].value,
        "대표자생년월일": sheet['F3'].value,
        "대표자거주지": sheet['G3'].value        
    }

def read_basicInfo(workbook):
    sheet = workbook['웰컴입력정보']
    date_of_meeting = sheet['C1'].value
    if isinstance(date_of_meeting, datetime):
        date_of_meeting = date_of_meeting.strftime('%Y년 %m월 %d일')
    return {
        "진행날짜": date_of_meeting,
        "관할등기소": sheet['C4'].value,
        "등록면허세": sheet['C5'].value,
        "지방교육세": sheet['C6'].value,
        "농어촌특별세": sheet['C7'].value,
        "세액합": (sheet['C5'].value) + (sheet['C6'].value) + (sheet['C7'].value),
        "등기수수료": (sheet['C8'].value),
        "납입처": (sheet['C9'].value)
    }





def format_number_with_commas(number):
    if isinstance(number, (int, float)):
        return f"{number:,}"
    return number

def replace_text_in_element(element, data_dict, apply_style=False):
    # _Cell 타입인지 확인하는 부분 추가
    if hasattr(element, 'text'):
        for key, value in data_dict.items():
            if isinstance(value, (int, float)):
                value = format_number_with_commas(value)
            elif value is None:
                value = ''

            # 텍스트에 키가 포함된 경우 대체
            if f"{{{{{key}}}}}" in element.text:
                element.text = element.text.replace(f"{{{{{key}}}}}", str(value))
                
                # 입력값에만 스타일 적용
                if apply_style and hasattr(element, 'paragraphs'):
                    for paragraph in element.paragraphs:
                        set_paragraph_style(paragraph)
    else:
        # 기본적인 단락이나 텍스트 객체를 처리
        for key, value in data_dict.items():
            if isinstance(value, (int, float)):
                value = format_number_with_commas(value)
            elif value is None:
                value = ''

            if f"{{{{{key}}}}}" in element.text:
                element.text = element.text.replace(f"{{{{{key}}}}}", str(value))
                
                if apply_style:
                    set_paragraph_style(element)


def remove_unused_markers(doc, markers):
    # 문서의 모든 단락과 표에서 남아 있는 마커를 빈 문자열로 대체하여 삭제
    for paragraph in doc.paragraphs:
        for marker in markers:
            if marker in paragraph.text:
                paragraph.text = paragraph.text.replace(marker, '')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for marker in markers:
                    if marker in cell.text:
                        cell.text = cell.text.replace(marker, '')

def set_paragraph_style(paragraph, font_name="맑은 고딕", font_size=10, alignment='left'):
    # 글씨체 및 글씨 크기 설정
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 한글 글꼴 설정
        run.font.size = Pt(font_size)

    # 정렬 설정
    if alignment == 'left':
        paragraph.alignment = 0  # 왼쪽 정렬
    elif alignment == 'center':
        paragraph.alignment = 1  # 가운데 정렬
    elif alignment == 'right':
        paragraph.alignment = 2  # 오른쪽 정렬

def generate_word_doc(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count, filepath):
    base_path = get_current_path()
    # 법인명 폴더 생성
    output_folder = os.path.join(base_path, '출력폴더(Word)', corp_info['법인명'])
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    template_base_path = get_template_path()
    template_path = os.path.join(template_base_path, 'doc_templates', '발기인회의사록_template.docx')
    output_filename = os.path.join(output_folder, f"발기인회의사록_{corp_info['법인명']}.docx")
    
    doc = Document(template_path)

    director_markers = [f"{{{{임원{i}}}}}" for i in range(3, 9)]

    for paragraph in doc.paragraphs:
        # apply_style=True로 설정하여 입력값에만 스타일 적용
        replace_text_in_element(paragraph, corp_info, apply_style=True)
        replace_text_in_element(paragraph, directors, apply_style=True)
        replace_text_in_element(paragraph, shareholders, apply_style=True)
        replace_text_in_element(paragraph, ceo_info, apply_style=True)
        replace_text_in_element(paragraph, basic_info, apply_style=True)
        replace_text_in_element(paragraph, {'주주수': shareholder_count}, apply_style=True)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_element(cell, corp_info, apply_style=True)
                replace_text_in_element(cell, directors, apply_style=True)
                replace_text_in_element(cell, shareholders, apply_style=True)
                replace_text_in_element(cell, ceo_info, apply_style=True)
                replace_text_in_element(cell, basic_info, apply_style=True)
                replace_text_in_element(cell, {'주주수': shareholder_count}, apply_style=True)

    remove_unused_markers(doc, director_markers)

    doc.save(output_filename)
    print(f"Document saved as {output_filename}")



def generate_word_doc_1(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count, filepath):
    base_path = get_current_path()
    # 법인명 폴더 생성
    output_folder = os.path.join(base_path, '출력폴더(Word)', corp_info['법인명'])
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    template_base_path = get_template_path()
    template_path = os.path.join(template_base_path, 'doc_templates', '정관_template.docx')
    output_filename = os.path.join(output_folder, f"정관_{corp_info['법인명']}.docx")
    
    
    doc = Document(template_path)

    conditional_texts = {
        '증자계획1': ("본 회사는 이사회 결의를 통하여 주주 및 주주 외의 제3자에게 신주를 발행할 수 있다." if corp_info['증자'] == '유' else "본 회사는 이사회 결의를 통하여 주주에게 신주를 발행할 수 있다."),
        '증자계획2': (
            "① 본 회사의 주주는 상법 제418조 제1항에 따라 그가 가진 주식 수에 비례하여 신주의 배정을 받을 수 있다. \n"
            "② 제1항의 규정에 불구하고 상법 제418조 제2항에 따라 본 회사의 경영상 목적을 달성하기 위하여 필요한 다음 각 호의 경우에는 주주 외의 제3자에게 신주를 배정할 수 있다.\n"
            "\t1. 신기술의 도입, 재무구조의 개선을 위하여 필요한 경우\n"
            "\t2. 자본시장과 금융투자업에 관한 법률 제165조의6에 의거하여 이사회의 결의를 통하여 일반공모 방식으로 신주를 발행하는 경우\n"
            "\t3. 자본시장과 금융투자업에 관한 법률 제165조의7의 규정에 의거하여 우리사주조합원에게 신주를 배정하는 경우\n"
            "\t4. 상법 또는 벤처기업육성에 관한 특별조치법에 의한 주식매수선택권의 행사로 인하여 신주를 발행하는 경우\n"
            "\t5. 긴급한 자금의 조달을 위하여 국내외 금융기관, 일반법인, 개인에게 신주를 발행하는 경우\n"
            "③ 본 회사가 벤처기업육성에 관한 특별조치법에 따라 벤처기업으로 확인받은 경우, 본 회사는 전략적 제휴를 위하여 주주총회 특별결의로 신주를 발행하여 다른 주식회사의 주요주주의 주식이나 주식회사인 다른 벤처기업의 주식과 교환할 수 있다. \n"
            "④	주주가 신주인수권을 포기 또는 상실하거나 신주배정에서 단주가 발생하는 경우에 그 처리방법은 이사회 결의로 정한다.\n"
            "⑤	신주에 대한 이익배당을 하는 경우, 신주를 발행한 시점이 속하는 사업연도의 직전 사업연도의 말에 발행된 것으로 보고 이익배당을 실시한다.\n" if corp_info['증자'] == '유' 
            else 
                "① 본 회사의 주주는 상법 제418조 제1항에 따라 그가 가진 주식 수에 비례하여 신주의 배정을 받을 수 있다.\n"
                "② 주주가 신주인수권을 포기 또는 상실하거나 신주배정에서 단주가 발생하는 경우에 그 처리방법은 이사회 결의로 정한다.\n"
                "③ 신주에 대한 이익배당을 하는 경우, 신주를 발행한 시점이 속하는 사업연도의 직전 사업연도의 말에 발행된 것으로 보고 이익배당을 실시한다\n"
        ),
        '중간배당':(
            "제48조의2 (중간배당)\n"
            "① 본 회사는 사업연도 중 1회에 한하여 이사회 결의로 일정한 날(이하 본 조에서 ‘기준일’이라 한다)을 정하여 그 날의 주주에 대하여 이익을 배당(이하 본 조에서 ‘중간배당’이라 한다)할 수 있다. 단, 당해 결산기의 재무상태표(대차대조표)상의 순자산액이 상법 제462조 제1항 각호의 금액의 합계액에 미치지 못할 우려가 있는 때는 중간배당을 할 수 없다.\n"
            "② 중간배당은 직전결산기의 재무상태표(대차대조표)상의 순자산액에서 다음 각호의 금액을 공제한 액을 한도로 한다.\n"
            "\t1. 직전 결산기의 자본금의 액\n"
            "\t2. 직전 결산기까지 적립된 자본준비금과 이익준비금의 합계액\n"
            "\t3. 직전 결산기의 정기주주총회에서 이익으로 배당하거나 지급하기로 정한 금액\n"
            "\t4. 중간배당에 따라 당해 결산기에 적립해야 할 이익준비금\n"
            "③ 본 회사가 사업연도개시일 이후 제1항의 기준일 이전에 신주를 발행한 경우(준비금의 자본전입, 주식배당, 전환사채의 전환청구, {{신주인수권부사채3}}주식매수선택권의 행사 등), 해당 신주는 직전사업연도말에 발행된 것으로 간주하여 중간배당을 실시한다.\n"
            "④제1항의 중간배당은 이사회 결의로 하며, 본 회사의 자본금 총액이 10억원 미만인 경우 이사회 결의사항은 주주총회 결의로 한다.\n"
        ),
        '신주인수권부사채1':(
            "제20조의3 (신주인수권부사채의 발행)\n"
            "① 본 회사가 이사회 결의로 신주인수권부사채를 발행하는 경우 다음 각호의 방식에 의한다.\n"
            "\t1. 주주에게 그가 가진 주식의 수에 따라서 신주인수권부사채를 배정하는 방식\n"
            "\t2. 사채의 액면총액이 100억원을 초과하지 않는 범위 내에서 신기술의 도입, 재무구조의 개선 등 회사의 경영상 목적을 달성하기 위하여 필요한 경우 특정한 자(이 회사의 주주를 포함한다)에게 사채를 배정하기 위하여 사채인수 청약을 할 기회를 부여하는 방식(단, 본 회사의 자본금 총액이 10억원 미만으로서 이사가 3인 미만인 경우 이사회 결의사항은 주주총회 특별결의로 한다)\n"
            "② 신주인수권부사채의 발행에 관하여 상법 제516조의2 제2항 각호의 사항은 이사회 결의로 정한다.\n"
            "③ 신주인수를 청구할 수 있는 금액은 사채의 액면총액을 초과하지 않는 범위내에서 이사회가 정한다.\n"
            "④ 신주인수권의 행사로 발행하는 주식은 보통주식으로 하고 그 발행가액은 액면금액 또는 그 이상의 가액으로 사채발행시 이사회가 정한다.\n"
            "⑦	신주인수권을 행사할 수 있는 기간은 당해 사채 발행일부터 그 상환기일까지 기간 내에서 이사회 결의로 정한다.\n"
            "⑧	다만 본 회사의 자본금 총액이 10억원 미만으로서 이사가 3인 미만인 경우로서 본 조에서 다르게 정하지 않은 경우 본 조의 이사회 결의사항은 주주총회 결의로 한다.\n" if corp_info['신주인수권부사채발행'] == '유' 
            else ""
        ),
        '신주인수권부사채2':("신주인수권부사채의 신주인수권 행사에 의하여 신주를 발행하는 경우, " if corp_info['신주인수권부사채발행'] == '유' else ""),
        '신주인수권부사채3':("신주인수권 행사, " if corp_info['신주인수권부사채발행'] == '유' else ""),
        '전환사채1':(
            "제20조의2 (전환사채의 발행) ① 본 회사가 이사회 결의로 전환사채를 발행하는 경우 다음 각호의 방식에 의한다.\n"
            "\t1. 주주에게 그가 가진 주식의 수에 따라서 전환사채를 배정하는 방식\n"
            "\t2. 사채의 액면총액이 100억원을 초과하지 않는 범위 내에서 신기술의 도입, 재무구조의 개선 등 회사의 경영상 목적을 달성하기 위하여 필요한 경우 특정한 자(이 회사의 주주를 포함한다)에게 사채를 배정하기 위하여 사채인수 청약을 할 기회를 부여하는 방식(단, 본 회사의 자본금 총액이 10억원 미만으로서 이사가 3인 미만인 경우 이사회 결의사항은 주주총회 특별결의로 한다)\n"
            "② 전환사채의 발행에 관하여 상법 제513조 제2항 각호의 사항은 이사회 결의로 정한다.\n"
            "③ 제1항의 전환사채에 있어서 이사회는 그 일부에 대하여만 전환권을 부여하는 조건으로도 이를 발행할 수 있다.\n"
            "④ 전환으로 인하여 발행하는 주식은 보통주식으로 하고 전환가액은 주식의 액면금액 또는 그 이상의 가액으로 사채발행시 이사회가 정한다.\n"
            "⑥ 전환을 청구할 수 있는 기간은 해당 사채의 발행일 후 1개월이 경과하는 날부터, 그 상환기일의 직전일까지로 하되, 이사회의 결의로 전환청구 기간을 조정할 수 있다.\n"
            "⑦ 주식으로 전환된 경우 회사는 전환 전에 지급시기가 도래한 이자에 대하여만 이자를 지급한다.\n"
            "⑧ 다만 본 회사의 자본금 총액이 10억원 미만으로서 이사가 3인 미만인 경우로서 본 조에서 다르게 정하지 않은 경우 본 조의 이사회 결의사항은 주주총회 결의로 한다.\n" if corp_info['전환사채'] == '유' else ""
        ),
        '전환사채2':("전환사채의 전환청구, " if corp_info['전환사채'] == '유' else ""),
    }

    # Replace text in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_element(paragraph, conditional_texts, apply_style=True)
        replace_text_in_element(paragraph, corp_info, apply_style=True)
        replace_text_in_element(paragraph, directors, apply_style=True)
        replace_text_in_element(paragraph, shareholders, apply_style=True)
        replace_text_in_element(paragraph, ceo_info, apply_style=True)
        replace_text_in_element(paragraph, basic_info, apply_style=True)
        replace_text_in_element(paragraph, {'주주수': shareholder_count}, apply_style=True)
        
        # 스타일 설정
        set_paragraph_style(paragraph, font_name="맑은 고딕", font_size=10, alignment='left')

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_element(cell, conditional_texts, apply_style=True)
                replace_text_in_element(cell, corp_info, apply_style=True)
                replace_text_in_element(cell, directors, apply_style=True)
                replace_text_in_element(cell, shareholders, apply_style=True)
                replace_text_in_element(cell, ceo_info, apply_style=True)
                replace_text_in_element(cell, basic_info, apply_style=True)
                replace_text_in_element(cell, {'주주수': shareholder_count}, apply_style=True)

                # 셀 내의 각 문단에도 스타일 설정
                for paragraph in cell.paragraphs:
                    set_paragraph_style(paragraph, font_name="맑은 고딕", font_size=10, alignment='left')

    doc.save(output_filename)
    print(f"Document saved as {output_filename}")


if __name__ == "__main__":
    window = Tk()
    settingUi(window)
    window.mainloop()