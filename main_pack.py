import os
import sys
import subprocess  # 폴더 열기 위해 subprocess 사용


from docx import Document
from datetime import datetime
from tkinter import Tk, Label, Button, Frame, Text, Scrollbar, RIGHT, Y, BOTH, END

# 전역 변수로 로그 Text 위젯을 참조할 수 있도록 설정
log_text_widget = None


def get_template_path():
    # doc_templates 폴더의 경로 반환 (패키징된 경우와 개발 중 경우 모두 지원)
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller로 패키징된 임시 폴더에서 경로를 가져옴
        return sys._MEIPASS
    return os.path.abspath(os.path.dirname(__file__))


def get_current_path():
    # 현재 실행 파일의 디렉토리를 반환 (입출력 폴더 경로 기반)
    return os.path.abspath(os.path.dirname(sys.argv[0]))

def settingUi(window):
    global log_text_widget  # 전역 변수로 설정하여 다른 함수에서 접근 가능

    window.title("법인설립서 자동작성 프로그램")  # 프로그램 타이틀

    screen_width = window.winfo_screenwidth()  # 화면 너비
    screen_height = window.winfo_screenheight()  # 화면 높이

    window.geometry("800x900+{}+{}".format(int(screen_width / 2 - 400), int(screen_height / 2 - 450)))  # 프로그램 크기 및 위치
    window.resizable(False, False)  # 프로그램 크기 조절 불가능

    # 프로그램 설명 프레임
    desc_frame = Frame(window, bg="#f0f0f0", padx=10, pady=10)
    desc_frame.pack(fill='x', pady=10)

    Label(desc_frame, text="기능 설명:", bg="#f0f0f0", anchor='w').pack(anchor='w')
    Label(desc_frame, text="- '입출력 폴더생성': 입력 및 출력 폴더를 생성합니다.", bg="#f0f0f0", anchor='w').pack(pady=2, padx=10, anchor='w')
    Label(desc_frame, text="- '법인서류 자동생성': 입력 폴더 내의 엑셀 파일을 읽고 문서를 생성합니다.", bg="#f0f0f0", anchor='w').pack(pady=2, padx=10, anchor='w')

    # 버튼 프레임
    button_frame = Frame(window)
    button_frame.pack(pady=20, fill='x')

    # 로그 프레임
    log_frame = Frame(window, padx=20, pady=15)
    log_frame.pack(pady=10, fill='both', expand=True)

    scrollbar = Scrollbar(log_frame)
    scrollbar.pack(side=RIGHT, fill=Y)

    log_text_widget = Text(log_frame, wrap='word', yscrollcommand=scrollbar.set)
    log_text_widget.pack(fill=BOTH, expand=True, padx=5, pady=5)

    scrollbar.config(command=log_text_widget.yview)

    # 버튼 추가
    Button(button_frame, text="입출력 폴더생성", width=20, height=2, command=createFolders).pack(side='left', expand=True, padx=10)
    Button(button_frame, text="법인서류 자동생성", width=20, height=2, command=checkInputExcel).pack(side='left', expand=True, padx=10)
    Button(button_frame, text="Templelet 폴더 열기", width=20, height=2, command=open_template_folder).pack(side='left', expand=True, padx=10)

def open_template_folder():
    # 템플릿 폴더 경로
    base_path = get_template_path()
    template_folder = os.path.join(base_path, 'doc_templates')

    # 템플릿 폴더가 존재하는지 확인하고 열기
    if os.path.exists(template_folder):
        if os.name == 'nt':  # Windows
            subprocess.Popen(f'explorer "{template_folder}"')
        elif os.name == 'posix':  # macOS 또는 Linux
            subprocess.Popen(['open', template_folder] if sys.platform == 'darwin' else ['xdg-open', template_folder])
    else:
        log_message("템플릿 폴더가 존재하지 않습니다", 'red')



def log_message(message, color='black'):
    if log_text_widget:
        tag_name = f"tag_{log_text_widget.index('end')}"  # 각 메시지에 고유한 태그 생성
        log_text_widget.insert(END, message + "\n", tag_name)
        log_text_widget.tag_config(tag_name, foreground=color)
        log_text_widget.see(END)

def createFolders():
    base_path = get_current_path()
    input_folder = os.path.join(base_path, '입력폴더(Excel)')
    output_folder = os.path.join(base_path, '출력폴더(Word)')


    checkFolders(input_folder, "입력")
    checkFolders(output_folder, "출력")

def checkFolders(folder_path, folder_name):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        log_message(f"{folder_name} 폴더가 생성되었습니다", 'green')
    else:
        log_message(f"{folder_name} 폴더가 이미 존재합니다", 'blue')

def checkInputExcel():
    base_path = get_current_path()
    input_folder_path = os.path.join(base_path, "입력폴더(Excel)")
    input_files = [f for f in os.listdir(input_folder_path) if f.endswith('.xlsx')]

    file_count = len(input_files)
    log_message(f"입력 폴더에 {file_count}개의 엑셀 파일이 발견되었습니다", 'blue')

    if not input_files:
        log_message("입력 폴더에 엑셀 파일이 존재하지 않습니다", 'red')
        return

    for idx, file_name in enumerate(input_files, start=1):
        log_message(f"({idx}/{file_count}) '{file_name}' 파일을 처리 중입니다...", 'blue')
        read_excel_data(os.path.join(input_folder_path, file_name))


def read_excel_data(filepath):
    from openpyxl import load_workbook
    workbook = load_workbook(filepath, data_only=True)
    corp_info, directors, shareholders, welcome_info = None, [], [], None

    try:
        corp_info = read_corpInfo(workbook)
        directors = read_directorIinfo(workbook)
        shareholders, shareholder_count = read_shareholderInfo(workbook)
        ceo_info = read_representativeInfo(workbook)
        basic_info = read_basicInfo(workbook)
        log_message(f"엑셀 파일 '{os.path.basename(filepath)}'의 정보를 불러왔습니다", 'blue')
        log_message(f"법인설립정보: {corp_info}")
        log_message(f"임원 정보: {directors}")
        log_message(f"주주 정보: {shareholders}")
        log_message(f"대표자 정보: {ceo_info}")
        log_message(f"기본 정보: {basic_info}")


        generate_word_doc(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count  ,filepath)
        generate_word_doc_1(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count  ,filepath)

    except KeyError as e:
        log_message(f"'{filepath}' 파일에서 시트를 찾을 수 없습니다: {e}", 'red')
    except Exception as e:
        log_message(f"'{filepath}' 파일 처리 중 에러 발생: {e}", 'red')
    finally:
        workbook.close()

def read_corpInfo(workbook):
    sheet = workbook['법인설립정보']
    return {
        "법인명": sheet['C3'].value,
        "영문법인명": sheet['C5'].value,
        "본점주소": sheet['C7'].value,
        "주당금액": sheet['C8'].value,
        "최대주식수": sheet['C9'].value,
        "발행주식수": sheet['C10'].value,
        "주식타입": sheet['C11'].value,
        "자본금": (sheet['C8'].value or 0) * (sheet['C10'].value or 0),
        "설립목적": sheet['C13'].value
    }

def read_directorIinfo(workbook):
    sheet = workbook['임원 주주 정보']
    directors = {}
    
    for i in range(3, 11):
        position = sheet[f'B{i}'].value
        nationality = sheet[f'C{i}'].value
        name = sheet[f'D{i}'].value
        english_name = sheet[f'E{i}'].value
        birth_date = sheet[f'F{i}'].value
        address = sheet[f'G{i}'].value

        if not position or not name:
            continue  # 직책이나 이름이 없으면 건너뜀
            

        if i == 3:
            if any(keyword in nationality.lower() for keyword in ["한국", "korean", "korea", "한국인"]):
                directors['대표피선자'] = f"대표이사 {name} ({birth_date})" + '\n'
                director_info = f"{position} {name} ({birth_date}) {address}" + '\n'
            else :
                directors['대표피선자'] = f"대표이사 {name} (영문:{english_name}, 국적:{nationality}, 생년월일:{birth_date})" + '\n'
                director_info = f"{position} {nationality} {name} {english_name} ({birth_date}) {address}" + '\n'
        else:
            if any(keyword in nationality.lower() for keyword in ["한국", "korean", "korea", "한국인"]):
                director_info = f"{position} {name} ({birth_date})" + '\n'
            else :
                director_info = f"{position} {nationality} {name} {english_name} ({birth_date})" + '\n'
        
        directors[f'임원{i - 2}'] = director_info + '\n'
    
    return directors


def read_shareholderInfo(workbook):
    sheet = workbook['임원 주주 정보']
    shareholders = {}
    shareholder_count = 0  # 주주 숫자 초기화

    for i in range(15, 23):
        row = [sheet[f'{chr(66 + j)}{i}'].value if sheet[f'{chr(66 + j)}{i}'].value is not None else '' for j in range(6)]
        
        # 모든 데이터가 비어 있는 경우 건너뜀
        if not any(row):
            continue
        
        shareholder_count += 1  # 주주가 발견될 때마다 증가
        shareholders[f'주주{shareholder_count}'] = row
    
    return shareholders, shareholder_count  # 주주 정보와 주주 숫자 반환



def read_representativeInfo(workbook):
    sheet = workbook['임원 주주 정보']

    return{
        "대표자직책": sheet['B3'].value,
        "대표자국적": sheet['C3'].value,
        "대표자명": sheet['D3'].value,
        "대표자영문명": sheet['E3'].value,
        "대표자생년월일": sheet['F3'].value,
        "대표자거주지": sheet['G3'].value        
    }

def read_basicInfo(workbook):
    sheet = workbook['웰컴입력정보']
    date_of_meeting = sheet['C1'].value
    if isinstance(date_of_meeting, datetime):
        date_of_meeting = date_of_meeting.strftime('%Y년 %m월 %d일')
    return {
        "진행날짜": date_of_meeting,
        "관할등기소": sheet['C4'].value,
        "등록면허세": sheet['C5'].value,
        "지방교육세": sheet['C6'].value,
        "농어촌특별세": sheet['C7'].value,
        "세액합": (sheet['C5'].value) + (sheet['C6'].value) + (sheet['C7'].value),
        "등기수수료": (sheet['C8'].value),
        "납입처": (sheet['C9'].value)
    }





def format_number_with_commas(number):
    if isinstance(number, (int, float)):
        return f"{number:,}"
    return number

def replace_text_in_element(element, data_dict):
    for key, value in data_dict.items():
        if isinstance(value, (int, float)):
            value = format_number_with_commas(value)
        elif value is None:
            value = ''  # 데이터가 없는 경우 빈 문자열로 대체

        # 마커가 텍스트에 포함되어 있을 경우 대체
        if f"{{{{{key}}}}}" in element.text:
            element.text = element.text.replace(f"{{{{{key}}}}}", str(value))

def remove_unused_markers(doc, markers):
    # 문서의 모든 단락과 표에서 남아 있는 마커를 빈 문자열로 대체하여 삭제
    for paragraph in doc.paragraphs:
        for marker in markers:
            if marker in paragraph.text:
                paragraph.text = paragraph.text.replace(marker, '')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for marker in markers:
                    if marker in cell.text:
                        cell.text = cell.text.replace(marker, '')

def generate_word_doc(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count, filepath):
    base_path = get_current_path()
    # 법인명 폴더 생성
    output_folder = os.path.join(base_path, '출력폴더(Word)', corp_info['법인명'])
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    template_base_path = get_template_path()
    template_path = os.path.join(template_base_path, 'doc_templates', '발기인회의사록_template.docx')
    output_filename = os.path.join(output_folder, f"발기인회의사록_{corp_info['법인명']}.docx")
    
    doc = Document(template_path)

    # 임원 마커 리스트 정의 (예: {{임원3}}, {{임원4}}, ...)
    director_markers = [f"{{{{임원{i}}}}}" for i in range(3, 9)]

    # Replace text in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_element(paragraph, corp_info)
        replace_text_in_element(paragraph, directors)
        replace_text_in_element(paragraph, shareholders)
        replace_text_in_element(paragraph, ceo_info)
        replace_text_in_element(paragraph, basic_info)
        replace_text_in_element(paragraph, {'주주수': shareholder_count})

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_element(cell, corp_info)
                replace_text_in_element(cell, directors)
                replace_text_in_element(cell, shareholders)
                replace_text_in_element(cell, ceo_info)
                replace_text_in_element(cell, basic_info)
                replace_text_in_element(cell, {'주주수': shareholder_count})

    # Unused marker 제거
    remove_unused_markers(doc, director_markers)

    # Save the modified document
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")


def generate_word_doc_1(corp_info, directors, shareholders, ceo_info, basic_info, shareholder_count, filepath):
    base_path = get_current_path()
    # 법인명 폴더 생성
    output_folder = os.path.join(base_path, '출력폴더(Word)', corp_info['법인명'])
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    template_base_path = get_template_path()
    template_path = os.path.join(template_base_path, 'doc_templates', '정관_template.docx')
    output_filename = os.path.join(output_folder, f"정관_{corp_info['법인명']}.docx")
    
    doc = Document(template_path)

    # Replace text in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_element(paragraph, corp_info)
        replace_text_in_element(paragraph, directors)
        replace_text_in_element(paragraph, shareholders)
        replace_text_in_element(paragraph, ceo_info)
        replace_text_in_element(paragraph, basic_info)
        replace_text_in_element(paragraph, {'주주수': shareholder_count})

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_element(cell, corp_info)
                replace_text_in_element(cell, directors)
                replace_text_in_element(cell, shareholders)
                replace_text_in_element(cell, ceo_info)
                replace_text_in_element(cell, basic_info)
                replace_text_in_element(cell, {'주주수': shareholder_count})

    # Save the modified document
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")


if __name__ == "__main__":
    window = Tk()
    settingUi(window)
    window.mainloop()